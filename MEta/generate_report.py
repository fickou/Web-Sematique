from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # Title Page
    document.add_heading('Rapport de Projet : Web Sémantique', 0)
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Groupes N° 5 et 6 : RDFS et OWL\n')
    run.bold = True
    p.add_run('Description sémantique des ressources de l’enseignement supérieur et de la recherche\n')
    
    document.add_page_break()

    # 1. Contexte et Problématique
    document.add_heading('1. Contexte et Problématique', level=1)
    
    document.add_heading('1.1 Contexte', level=2)
    document.add_paragraph(
        "Les établissements d’enseignement supérieur et de recherche (universités, écoles, laboratoires) "
        "manipulent une grande variété de ressources : formations, personnels (enseignants, chercheurs, étudiants), "
        "publications, infrastructures, etc. Actuellement, ces ressources sont gérées dans des systèmes d'information "
        "hétérogènes et cloisonnés (Scolarité, RH, Gestion de la recherche)."
    )

    document.add_heading('1.2 Problématique', level=2)
    document.add_paragraph(
        "Cette fragmentation des données pose un problème majeur d'interopérabilité. Il est difficile d'avoir une "
        "vue unifiée des ressources (par exemple, lier facilement un projet de recherche aux formations doctorales "
        "et aux équipements utilisés). Le Web Sémantique offre des standards (RDF, RDFS, OWL) pour décrire ces "
        "données de manière normée et interconnectable."
    )

    document.add_heading('1.3 Objectifs', level=2)
    document.add_paragraph(
        "L'objectif de ce projet est de concevoir et implémenter une ontologie modulaire permettant de décrire "
        "et de lier sémantiquement les ressources académiques. Nous avons utilisé RDFS pour la structure de base "
        "et OWL pour enrichir la sémantique (contraintes, inférences)."
    )

    # 2. Travail Réalisé
    document.add_heading('2. Travail Réalisé', level=1)

    document.add_heading('2.1 Choix de Modélisation (RDFS)', level=2)
    document.add_paragraph(
        "Nous avons structuré notre vocabulaire autour de plusieurs axes principaux :"
    )
    document.add_paragraph("• Ressources Humaines : Hiérarchie de classes avec concept racine 'Personne', spécialisée en 'Enseignant', 'Etudiant', 'PersonnelAdministratif'.", style='List Bullet')
    document.add_paragraph("• Structure Académique : 'Universite', 'Faculte', 'Departement', 'Laboratoire'.", style='List Bullet')
    document.add_paragraph("• Formation : 'Formation' (Licence, Master, Doctorat) et 'UniteEnseignement' (UE).", style='List Bullet')
    document.add_paragraph("• Propriétés : Relationnelles (enseigne, inscritDans, membreDe) et Littérales (nom, email, crédits).", style='List Bullet')
    
    document.add_paragraph(
        "Nous avons aligné notre vocabulaire avec des standards existants : FOAF (pour les personnes) et Dublin Core (pour les métadonnées)."
    )

    document.add_heading('2.2 Contraintes OWL Utilisées', level=2)
    document.add_paragraph(
        "Pour garantir la cohérence et permettre des inférences, nous avons utilisé OWL :"
    )
    document.add_paragraph("• Cardinalités : Un étudiant est inscrit dans au moins une formation (minCardinality 1). Une formation contient au moins une UE.", style='List Bullet')
    document.add_paragraph("• Disjonction : Les classes 'Enseignant' et 'Etudiant' ont été déclarées disjointes (sauf exception doctorant enseignant, gérée par cas particulier si nécessaire).", style='List Bullet')
    document.add_paragraph("• Intersection et classes définies : Définition formelle de l'Enseignant-Chercheur (voir section inférences).", style='List Bullet')

    # 3. Exemples d'Inférences
    document.add_heading('3. Exemples d’Inférences', level=1)
    
    document.add_paragraph(
        "L'un des points forts de notre modélisation est la capacité du système à déduire de nouvelles informations "
        "non présentes explicitement dans les données brutes."
    )

    document.add_heading('3.1 Classification Automatique (Enseignant-Chercheur)', level=2)
    document.add_paragraph(
        "Problème : Dans les données brutes, un individu comme 'Prof_Ndiaye' est déclaré comme 'Enseignant' "
        "et possède une relation 'membreDe' vers un 'Laboratoire', mais n'est pas explicitement typé 'EnseignantChercheur'."
    )
    document.add_paragraph(
        "Solution OWL : Nous avons défini la classe 'EnseignantChercheur' comme une classe équivalente à l'intersection "
        "de la classe 'Enseignant' et de la restriction 'membreDe some Laboratoire'."
    )
    document.add_paragraph(
        "Résultat : Le raisonneur infère automatiquement que 'Prof_Ndiaye' est un 'EnseignantChercheur'. "
        "Sans cette inférence, une requête SPARQL sur les enseignants-chercheurs ne retournait aucun résultat."
    )

    document.add_heading('3.2 Inférence de Types (RDFS)', level=2)
    document.add_paragraph(
        "Grâce à la hiérarchie 'rdfs:subClassOf', toute instance d'Etudiant ou d'Enseignant est automatiquement "
        "inférée comme étant une 'Personne'. Cela permet d'interroger globalement l'annuaire via la classe racine 'Personne'."
    )

    # 4. Perspectives
    document.add_heading('4. Perspectives', level=1)
    document.add_paragraph(
        "Ce travail pose les bases d'un système d'information sémantique universitaire. Les perspectives d'évolution sont nombreuses :"
    )
    document.add_paragraph("• Interconnexion (LOD) : Lier nos données avec celles d'autres universités ou référentiels (DBpedia, Wikidata) pour enrichir les profils chercheurs.", style='List Bullet')
    document.add_paragraph("• Moteur de Recherche Facetté : Créer une interface permettant de filtrer les ressources par facettes (par UFR, par domaine de recherche, par type de formation).", style='List Bullet')
    document.add_paragraph("• Système de Recommandation : Suggérer des UE optionnelles aux étudiants en fonction de leur parcours ou des projets de recherche liés.", style='List Bullet')

    document.save('Rapport_Projet_Web_Semantique.docx')
    print("Rapport généré avec succès : Rapport_Projet_Web_Semantique.docx")

if __name__ == "__main__":
    create_report()
