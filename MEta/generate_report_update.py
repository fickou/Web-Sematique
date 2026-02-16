import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_robust(doc, text, level=1):
    # Try typical style names
    style_names = [f"Heading {level}", f"Titre {level}"]
    
    success = False
    for name in style_names:
        try:
            doc.add_heading(text, level=level)
            success = True
            break
        except KeyError:
            continue
            
    if not success:
        # Fallback to manual formatting
        print(f"Heading style level {level} not found. Using fallback formatting.")
        p = doc.add_paragraph(text)
        run = p.runs[0]
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(46, 116, 181) # Blueish
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(46, 116, 181)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(31, 77, 120)

def add_caption_robust(doc, text):
    try:
        doc.add_paragraph(text, style='Caption')
    except KeyError:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.italic = True
        run.font.size = Pt(9)

def main():
    base_dir = "/home/mg4/Documents/VS code Master 2 GDIL 1er smtre/Web Semantique/Web-Sematique"
    report_path = os.path.join(base_dir, "Rapport_Projet_Web_Semantique.docx")
    output_path = os.path.join(base_dir, "Rapport_Projet_Web_Semantique_Updated.docx")
    
    # Image paths
    graph_img = os.path.join(base_dir, "graphe de donnees.png")
    graph_zoom_img = os.path.join(base_dir, "graphe de donnees_zoomer.png")
    rq_dir = os.path.join(base_dir, "RQ&resultat")

    doc = Document(report_path)

    # Find "2-Travail Réalisé" and insert after it
    target_para = None
    for para in doc.paragraphs:
        if "2-Travail Réalisé" in para.text or "2- Travail Réalisé" in para.text or "2 - Travail Réalisé" in para.text:
            target_para = para
            break
            
    if target_para:
        print("Found section '2-Travail Réalisé'. Appending content after it.")
        pass
    else:
        print("Section '2-Travail Réalisé' not found strictly. Appending to end.")
        add_heading_robust(doc, "2-Travail Réalisé", level=1)

    # SECTION: GRAPH ANALYSIS
    add_heading_robust(doc, "2.1 Analyse du Graphe de Données", level=2)
    
    p = doc.add_paragraph("Le graphe de données ci-dessous représente les relations sémantiques entre les différentes entités modélisées (Enseignants, Étudiants, Formations, etc.).")
    
    if os.path.exists(graph_img):
        doc.add_picture(graph_img, width=Inches(6))
        last_p = doc.paragraphs[-1] 
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption_robust(doc, "Figure 1: Graphe de données global")

    if os.path.exists(graph_zoom_img):
        doc.add_picture(graph_zoom_img, width=Inches(6))
        last_p = doc.paragraphs[-1] 
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption_robust(doc, "Figure 2: Graphe de données (vue zoomée)")

    doc.add_paragraph("Le graphe permet de visualiser la connectivité des données instanciées dans les fichiers Turtle (UGB et UASZ). On y observe les nœuds représentant les individus (ex: Prof_Diop, Etu_Fall) et les arcs représentant les propriétés (ex: enseigne, inscritDans).")

    # SECTION: QUERY ANALYSIS
    add_heading_robust(doc, "2.2 Exécution des Requêtes et Analyse des Résultats", level=2)
    
    queries = [
        {
            "id": 1,
            "title": "Requête 1 : Lister les enseignants et leurs UE",
            "desc": "Cette requête récupère les enseignants et les unités d'enseignement qu'ils dispensent.",
            "img_rq": "Rq1.png",
            "img_res": "resultat1.png",
            "comment": "Résultat obtenu sans inférence particulière. Les relations sont explicites dans les données."
        },
        {
            "id": 2,
            "title": "Requête 2 : Étudiants par Formation",
            "desc": "Affichage des étudiants inscrits dans chaque formation.",
            "img_rq": "Rq2.png",
            "img_res": "resultat2.png",
            "comment": "Liste des étudiants liés aux formations via la propriété 'ens:inscritDans'."
        },
        {
            "id": 3,
            "title": "Requête 3 : Recherche des Enseignants-Chercheurs",
            "desc": "Identification des individus de type EnseignantChercheur.",
            "img_rq": "Rq3.png",
            "img_res": "resultat3.png",
            "img_inf": "resultat_inf_3.png", 
            "comment_no_inf": "Sans inférence : Aucun résultat. Dans le fichier Turtle, personne n'est explicitement déclaré comme 'EnseignantChercheur'.",
            "comment_inf": "Avec inférence (OWLRL) : Le raisonneur déduit que les individus qui sont à la fois 'Enseignant' et membres d'un 'Laboratoire' (intersection de classes) sont des 'EnseignantChercheur'. Par exemple, Moussa Diop est classé automatiquement."
        },
        {
            "id": 4,
            "title": "Requête 4 : Doctorants et Directeurs de Thèse",
            "desc": "Liste des doctorants, leurs directeurs et écoles doctorales.",
            "img_rq": "Rq4.png",
            "img_res": "resultat4.png",
            "comment": "Relation directe exploitant les propriétés 'ens:dirigePar' et 'ens:inscritDansED'."
        },
        {
            "id": 5,
            "title": "Requête 5 : Lister toutes les Personnes",
            "desc": "Récupération de toutes les instances de la classe Personne.",
            "img_rq": "Rq5.png",
            "img_res": "resultat5.png",
            "img_inf": "resultat_inf_5.png", 
            "comment_no_inf": "Sans inférence : Aucun résultat (ou partiel) si la classe 'Personne' n'est pas instanciée directement.",
            "comment_inf": "Avec inférence : Le raisonneur utilise la hiérarchie des classes (rdfs:subClassOf). Puisque 'Enseignant', 'Etudiant', etc. sont des sous-classes de 'Personne', toutes leurs instances sont inférées comme étant des 'Personne'."
        }
    ]

    for q in queries:
        add_heading_robust(doc, f"{q['title']}", level=3)
        doc.add_paragraph(q['desc'])
        
        # Query Image
        rq_path = os.path.join(rq_dir, q['img_rq'])
        if os.path.exists(rq_path):
            doc.add_picture(rq_path, width=Inches(5))
            add_caption_robust(doc, f"Requête SPARQL {q['id']}")
        
        # Result Image (No Inference)
        res_path = os.path.join(rq_dir, q['img_res'])
        if os.path.exists(res_path):
            doc.add_paragraph("Résultat (Sans Inférence) :")
            doc.add_picture(res_path, width=Inches(5))
        
        # Inference Result Image (if applicable)
        if 'img_inf' in q:
            inf_path = os.path.join(rq_dir, q['img_inf'])
            if os.path.exists(inf_path):
                doc.add_paragraph("Résultat (Avec Inférence) :")
                doc.add_picture(inf_path, width=Inches(5))
                
                # Inference Commentary
                try:
                    p = doc.add_paragraph("Analyse :")
                    p.runs[0].bold=True
                    p.style = "Heading 4"
                except KeyError:
                    pass

                doc.add_paragraph(f"- {q['comment_no_inf']}")
                doc.add_paragraph(f"- {q['comment_inf']}")
            else:
                 doc.add_paragraph("Analyse : " + q['comment'])
        else:
             doc.add_paragraph("Analyse : " + q['comment'])

        doc.add_paragraph("_" * 50) # Separator

    doc.save(output_path)
    print(f"Report saved to {output_path}")

if __name__ == "__main__":
    main()
